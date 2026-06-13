"""Resume storage, text extraction, and OpenAI-powered tailoring per job."""
import io
import json
import logging
import os
import re
import threading
from datetime import datetime, timezone
from pathlib import Path

import requests

from . import db

log = logging.getLogger("jobfinder.resume")

DATA_DIR = db.DATA_HOME / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)

OPENAI_URL = "https://api.openai.com/v1/chat/completions"
DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

_lock = threading.Lock()

SCHEMA = """
CREATE TABLE IF NOT EXISTS tailored (
    job_id     TEXT PRIMARY KEY,
    content    TEXT NOT NULL,      -- tailored resume (markdown)
    keywords   TEXT DEFAULT '',    -- comma separated ATS keywords
    changes    TEXT DEFAULT '',    -- summary of what was changed
    model      TEXT DEFAULT '',
    created_at TEXT NOT NULL
);
"""


def init():
    with _lock, db.get_conn() as conn:
        conn.executescript(SCHEMA)


# ------------------------------------------------------------- resume store

def extract_text(filename: str, blob: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(blob))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if ext == ".docx":
        import docx
        document = docx.Document(io.BytesIO(blob))
        return "\n".join(p.text for p in document.paragraphs).strip()
    if ext in (".txt", ".md"):
        return blob.decode("utf-8", errors="replace").strip()
    raise ValueError(f"Unsupported file type {ext!r}; use PDF, DOCX, TXT or MD")


def save_resume(filename: str, blob: bytes) -> dict:
    text = extract_text(filename, blob)
    if len(text) < 200:
        raise ValueError("Couldn't extract enough text from that file — "
                         "is it a scanned image? Try a text-based PDF or DOCX.")
    (DATA_DIR / "resume_original" ).with_suffix(Path(filename).suffix.lower()).write_bytes(blob)
    (DATA_DIR / "resume_text.txt").write_text(text, encoding="utf-8")
    meta = {"filename": filename,
            "uploaded_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S"),
            "chars": len(text)}
    db.set_meta("resume", json.dumps(meta))
    return meta


def get_resume_meta() -> dict | None:
    raw = db.get_meta("resume")
    return json.loads(raw) if raw else None


def get_resume_text() -> str | None:
    p = DATA_DIR / "resume_text.txt"
    return p.read_text(encoding="utf-8") if p.exists() else None


# --------------------------------------------------------------- OpenAI key

def get_api_key() -> str | None:
    return os.environ.get("OPENAI_API_KEY") or db.get_meta("openai_key")


def set_api_key(key: str):
    db.set_meta("openai_key", key.strip())


# ----------------------------------------------------------------- tailoring

PROMPT = """You are an expert resume writer and ATS optimization specialist.

Below is a candidate's resume and a job posting. Rewrite the resume tailored to
this specific job. Rules:
- NEVER invent experience, employers, dates, degrees, or certifications.
- Reorder and rephrase existing content to emphasize what this job values.
- Naturally weave in the job's important keywords/skills where the candidate
  genuinely has them (for ATS matching).
- Keep it to a similar length as the original. Use clean Markdown:
  # name, ## section headings, - bullet points.
- Quantify achievements where the original already provides numbers.

Return strict JSON with these keys:
  "resume_md": the full tailored resume in Markdown,
  "keywords": array of the most important ATS keywords from the posting,
  "changes": array of short strings describing each change you made.

JOB POSTING
Title: {title}
Company: {company}
Required skills detected: {skills}
Description:
{description}

CANDIDATE'S RESUME
{resume}
"""


def tailor(job_id: str) -> dict:
    api_key = get_api_key()
    if not api_key:
        raise ValueError("No OpenAI API key set. Add it in Settings (or export OPENAI_API_KEY).")
    resume_text = get_resume_text()
    if not resume_text:
        raise ValueError("No resume uploaded yet. Upload your resume first.")

    with db.get_conn() as conn:
        job = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        raise ValueError("Job not found")
    job = dict(job)

    prompt = PROMPT.format(
        title=job["title"], company=job["company"],
        skills=job["skills"] or "n/a",
        description=job["description"][:4000],
        resume=resume_text[:12000],
    )
    body = {
        "model": DEFAULT_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.4,
        "response_format": {"type": "json_object"},
    }
    r = requests.post(OPENAI_URL, json=body, timeout=120,
                      headers={"Authorization": f"Bearer {api_key}"})
    if r.status_code == 401:
        raise ValueError("OpenAI rejected the API key (401). Check it in Settings.")
    r.raise_for_status()
    content = r.json()["choices"][0]["message"]["content"]
    try:
        parsed = json.loads(content)
        resume_md = parsed["resume_md"]
        keywords = parsed.get("keywords", [])
        changes = parsed.get("changes", [])
    except (json.JSONDecodeError, KeyError):
        resume_md, keywords, changes = content, [], []

    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S")
    with _lock, db.get_conn() as conn:
        conn.execute(
            """INSERT INTO tailored (job_id, content, keywords, changes, model, created_at)
               VALUES (?,?,?,?,?,?)
               ON CONFLICT(job_id) DO UPDATE SET content=excluded.content,
                 keywords=excluded.keywords, changes=excluded.changes,
                 model=excluded.model, created_at=excluded.created_at""",
            (job_id, resume_md, ",".join(keywords), json.dumps(changes),
             DEFAULT_MODEL, now),
        )
    return {"job_id": job_id, "resume_md": resume_md, "keywords": keywords,
            "changes": changes, "created_at": now}


def get_tailored(job_id: str) -> dict | None:
    with db.get_conn() as conn:
        row = conn.execute("SELECT * FROM tailored WHERE job_id = ?", (job_id,)).fetchone()
    if not row:
        return None
    d = dict(row)
    d["keywords"] = [k for k in d["keywords"].split(",") if k]
    try:
        d["changes"] = json.loads(d["changes"])
    except json.JSONDecodeError:
        d["changes"] = []
    d["resume_md"] = d.pop("content")
    return d


# ------------------------------------------------------------ docx rendering

def markdown_to_docx(md: str) -> bytes:
    """Render the tailored Markdown resume as a .docx file."""
    import docx
    from docx.shared import Pt

    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    bold_re = re.compile(r"\*\*(.+?)\*\*")

    def add_runs(par, text):
        pos = 0
        for m in bold_re.finditer(text):
            if m.start() > pos:
                par.add_run(text[pos:m.start()])
            par.add_run(m.group(1)).bold = True
            pos = m.end()
        if pos < len(text):
            par.add_run(text[pos:])

    for line in md.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            document.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            document.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            document.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ")):
            add_runs(document.add_paragraph(style="List Bullet"), s[2:])
        elif s.startswith("---"):
            continue
        else:
            add_runs(document.add_paragraph(), s)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
